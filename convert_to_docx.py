"""
Convert PROJECT_REPORT.md to DOCX format
Run: pip install python-docx markdown
     python convert_to_docx.py
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_project_report_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('3D Tactical Combat Arena', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Project Title'
    table.rows[0].cells[1].text = '3D Tactical Combat Arena - AI vs AI'
    table.rows[1].cells[0].text = 'Technology Stack'
    table.rows[1].cells[1].text = 'Python 3.10+, Ursina Engine, NumPy, Scikit-Fuzzy'
    table.rows[2].cells[0].text = 'Project Type'
    table.rows[2].cells[1].text = 'Turn-based Tactical AI Game'
    
    doc.add_paragraph()
    doc.add_paragraph('This project implements a production-grade 3D tactical combat game featuring two competing AI agents: Minimax with Alpha-Beta Pruning (Red Team) vs Fuzzy Logic Controller (Blue Team).')
    
    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Project Objective', level=2)
    doc.add_paragraph('To develop a turn-based tactical combat game that demonstrates the comparison between two fundamentally different AI approaches:')
    doc.add_paragraph('• Adversarial Search (Minimax) - Exhaustive game tree exploration', style='List Bullet')
    doc.add_paragraph('• Fuzzy Inference (Fuzzy Logic) - Human-like rule-based reasoning', style='List Bullet')
    
    doc.add_heading('1.2 Key Features', level=2)
    doc.add_paragraph('• Beautiful 3D graphics using Ursina engine', style='List Bullet')
    doc.add_paragraph('• 5 unique unit types with special abilities', style='List Bullet')
    doc.add_paragraph('• Dynamic terrain system affecting combat', style='List Bullet')
    doc.add_paragraph('• Real-time AI decision visualization', style='List Bullet')
    doc.add_paragraph('• Tournament mode for statistical analysis', style='List Bullet')
    
    # System Architecture
    doc.add_heading('2. System Architecture', level=1)
    doc.add_heading('2.1 Project Structure', level=2)
    
    structure = """PythonAdvanced3DGame/
├── main.py              # Entry point
├── ai/                  # AI agents
│   ├── minimax_agent.py # Minimax AI
│   ├── fuzzy_agent.py   # Fuzzy Logic AI
│   └── evaluation.py    # Heuristic evaluation
├── core/                # Game logic
│   ├── game_state.py    # State management
│   ├── battlefield.py   # Terrain system
│   └── unit.py          # Unit classes
├── graphics/            # 3D rendering
└── config/              # Configuration"""
    
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('2.2 Technology Stack', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Technology', 'Purpose']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    tech_data = [
        ('3D Engine', 'Ursina 5.0+', 'Graphics rendering'),
        ('Numerical', 'NumPy 1.24+', 'Array operations'),
        ('Fuzzy Logic', 'Scikit-Fuzzy 0.4.2', 'Fuzzy inference'),
        ('Validation', 'Pydantic 2.0+', 'Data models'),
        ('Testing', 'Pytest 7.0+', 'Unit tests'),
    ]
    for i, (comp, tech, purpose) in enumerate(tech_data, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = tech
        table.rows[i].cells[2].text = purpose
    
    # Game Mechanics
    doc.add_heading('3. Game Mechanics', level=1)
    doc.add_heading('3.1 Unit Types', level=2)
    
    table = doc.add_table(rows=6, cols=7)
    table.style = 'Table Grid'
    unit_headers = ['Unit', 'HP', 'Attack', 'Defense', 'Range', 'Move', 'Role']
    for i, h in enumerate(unit_headers):
        table.rows[0].cells[i].text = h
    
    units = [
        ('Warrior', '150', '30', '20', '1', '2', 'Tank/Frontline'),
        ('Archer', '80', '35', '8', '5', '3', 'Ranged DPS'),
        ('Mage', '70', '45', '5', '4', '2', 'Area Damage'),
        ('Knight', '120', '35', '15', '1', '4', 'Mobile Fighter'),
        ('Healer', '60', '15', '10', '3', '3', 'Support'),
    ]
    for i, unit in enumerate(units, 1):
        for j, val in enumerate(unit):
            table.rows[i].cells[j].text = val
    
    doc.add_heading('3.2 Terrain System', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    terrain_headers = ['Terrain', 'Defense', 'Attack', 'Movement']
    for i, h in enumerate(terrain_headers):
        table.rows[0].cells[i].text = h
    
    terrains = [
        ('Plain', '+0', '+0', '1'),
        ('Forest', '+2', '+0', '2'),
        ('Hill', '+1', '+1', '2'),
        ('Ruins', '+3', '+0', '1'),
        ('Water', 'N/A', 'N/A', 'Impassable'),
    ]
    for i, t in enumerate(terrains, 1):
        for j, val in enumerate(t):
            table.rows[i].cells[j].text = val
    
    doc.add_heading('3.3 Combat Formula', level=2)
    combat = """Raw Damage = Attacker's Attack + Terrain Attack Bonus
Defense = Target's Defense + Terrain Defense Bonus
Final Damage = max(1, Raw Damage - Defense)
Critical Hits = 1.5x damage"""
    p = doc.add_paragraph()
    run = p.add_run(combat)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # AI Implementation
    doc.add_heading('4. AI Implementation', level=1)
    
    doc.add_heading('4.1 Minimax Agent (Red Team)', level=2)
    doc.add_paragraph('Algorithm: Minimax with Alpha-Beta Pruning')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Alpha-Beta pruning for efficiency', style='List Bullet')
    doc.add_paragraph('• Transposition tables for caching', style='List Bullet')
    doc.add_paragraph('• Iterative deepening for time management', style='List Bullet')
    doc.add_paragraph('• Move ordering for better pruning', style='List Bullet')
    
    doc.add_paragraph('Decision Process Example:')
    decision = """Current: Red turn
├── Move Warrior forward → Score: 45
├── Attack with Archer → Score: 60 ⭐ BEST!
└── Wait → Score: 20

Decision: Attack with Archer (highest score)"""
    p = doc.add_paragraph()
    run = p.add_run(decision)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('4.2 Fuzzy Logic Agent (Blue Team)', level=2)
    doc.add_paragraph('Algorithm: Mamdani-style Fuzzy Inference System')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Threat assessment system', style='List Bullet')
    doc.add_paragraph('• Aggression level controller', style='List Bullet')
    doc.add_paragraph('• Target prioritization', style='List Bullet')
    doc.add_paragraph('• Human-readable rules', style='List Bullet')
    
    doc.add_paragraph('Fuzzy Rules:')
    rules = """IF HP is LOW AND enemies are CLOSE → Be DEFENSIVE
IF enemy HP is CRITICAL AND can_kill → PRIORITY TARGET
IF winning AND healthy → Be AGGRESSIVE"""
    p = doc.add_paragraph()
    run = p.add_run(rules)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # State Evaluation
    doc.add_heading('5. State Evaluation', level=1)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    eval_headers = ['Factor', 'Weight', 'Description']
    for i, h in enumerate(eval_headers):
        table.rows[0].cells[i].text = h
    
    factors = [
        ('HP Difference', '1.0', 'Total HP comparison'),
        ('Unit Count', '50.0', 'Surviving units'),
        ('Position', '0.5', 'Terrain & center control'),
        ('Threats', '0.3', 'Attack opportunities'),
        ('Mobility', '0.2', 'Movement options'),
        ('Terrain', '0.15', 'Terrain advantage'),
        ('Formation', '0.1', 'Unit coordination'),
    ]
    for i, f in enumerate(factors, 1):
        for j, val in enumerate(f):
            table.rows[i].cells[j].text = val
    
    # Battle Visualization
    doc.add_heading('6. Sample Battle Visualization', level=1)
    
    doc.add_heading('6.1 Initial Setup', level=2)
    initial = """┌───┬───┬───┬───┬───┬───┬───┬───┬───┬───┬───┬───┐
│🔴W│   │   │   │   │   │   │   │   │   │   │🔵K│
├───┼───┼───┼───┼───┼───┼───┼───┼───┼───┼───┼───┤
│🔴A│   │   │🌲 │   │   │   │   │🌲 │   │   │🔵A│
├───┼───┼───┼───┼───┼───┼───┼───┼───┼───┼───┼───┤
│🔴M│   │   │   │   │💧│💧│   │   │   │   │🔵M│
├───┼───┼───┼───┼───┼───┼───┼───┼───┼───┼───┼───┤
│🔴H│   │   │   │   │   │   │   │   │   │   │🔵H│
└───┴───┴───┴───┴───┴───┴───┴───┴───┴───┴───┴───┘
W=Warrior, A=Archer, M=Mage, H=Healer, K=Knight"""
    p = doc.add_paragraph()
    run = p.add_run(initial)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    
    doc.add_heading('6.2 Kill Sequence', level=2)
    kill = """🔴 Warrior attacks 🔵 Knight! → Damage: 15 → HP: 45→30
🔴 Archer follows up! → Damage: 20 → HP: 30→10
🔴 Mage casts Fireball! → HP: 10→0 ☠️

💀 BLUE KNIGHT ELIMINATED!"""
    p = doc.add_paragraph()
    run = p.add_run(kill)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # Special Abilities
    doc.add_heading('7. Special Abilities', level=1)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    ability_headers = ['Unit', 'Ability', 'Effect', 'Cooldown']
    for i, h in enumerate(ability_headers):
        table.rows[0].cells[i].text = h
    
    abilities = [
        ('Warrior', 'Shield Wall', '+5 Defense, 2 turns', '3 turns'),
        ('Warrior', 'Charge', 'Move + Attack', '4 turns'),
        ('Archer', 'Snipe', '+3 Range, +50% damage', '4 turns'),
        ('Archer', 'Overwatch', 'Counter-attack', '5 turns'),
        ('Mage', 'Fireball', '3×3 area damage', '4 turns'),
        ('Knight', 'Charge', 'Move + Attack', '3 turns'),
        ('Healer', 'Heal', 'Restore 40 HP', '2 turns'),
    ]
    for i, a in enumerate(abilities, 1):
        for j, val in enumerate(a):
            table.rows[i].cells[j].text = val
    
    # Conclusion
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph('This project successfully demonstrates:')
    doc.add_paragraph('1. AI Algorithm Comparison: Minimax (systematic) vs Fuzzy Logic (intuitive)', style='List Number')
    doc.add_paragraph('2. Production-Grade Architecture: Modular, testable, configurable', style='List Number')
    doc.add_paragraph('3. Rich Game Mechanics: Multiple units, terrain effects, special abilities', style='List Number')
    doc.add_paragraph('4. Visual Presentation: 3D graphics with real-time AI visualization', style='List Number')
    
    doc.add_heading('Future Enhancements', level=2)
    doc.add_paragraph('• Reinforcement learning agent', style='List Bullet')
    doc.add_paragraph('• Multiplayer support', style='List Bullet')
    doc.add_paragraph('• Additional unit types', style='List Bullet')
    doc.add_paragraph('• Network play capability', style='List Bullet')
    
    # Save
    doc.save('PROJECT_REPORT.docx')
    print("✅ PROJECT_REPORT.docx created successfully!")

if __name__ == '__main__':
    create_project_report_docx()
